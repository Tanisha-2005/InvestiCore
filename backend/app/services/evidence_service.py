import hashlib
import io
import email
from typing import Dict, Any, Tuple

# Optional forensic libraries with graceful fallback
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.storage import storage_service

class EvidenceService:
    def calculate_hashes(self, file_bytes: bytes) -> Tuple[str, str, str]:
        md5_hash = hashlib.md5(file_bytes).hexdigest()
        sha1_hash = hashlib.sha1(file_bytes).hexdigest()
        sha256_hash = hashlib.sha256(file_bytes).hexdigest()
        return md5_hash, sha1_hash, sha256_hash

    def extract_text(self, file_bytes: bytes, file_name: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        extracted_text = ""
        metadata = {}

        try:
            filename_lower = file_name.lower()

            # Image evidence OCR
            if file_type == "image" or any(filename_lower.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]):
                if OCR_AVAILABLE:
                    try:
                        img = Image.open(io.BytesIO(file_bytes))
                        extracted_text = pytesseract.image_to_string(img)
                        metadata = {"image_size": img.size, "image_format": img.format}
                    except Exception as e:
                        extracted_text = f"[OCR Image Processing Error: {str(e)}]"
                else:
                    extracted_text = "[Tesseract OCR module not installed on host. Run via Docker for full OCR.]"

            # PDF Evidence
            elif file_type == "pdf" or filename_lower.endswith(".pdf"):
                if PDF_AVAILABLE:
                    try:
                        doc = fitz.open(stream=file_bytes, filetype="pdf")
                        text_runs = [page.get_text() for page in doc]
                        extracted_text = "\n".join(text_runs)
                        metadata = {"pdf_pages": len(doc), "encrypted": doc.is_encrypted}
                    except Exception as e:
                        extracted_text = f"[PDF Parsing Error: {str(e)}]"
                else:
                    extracted_text = file_bytes.decode("utf-8", errors="ignore")

            # Office Document (.docx)
            elif file_type == "office" or filename_lower.endswith(".docx"):
                if DOCX_AVAILABLE:
                    try:
                        doc = docx.Document(io.BytesIO(file_bytes))
                        extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text])
                        metadata = {"docx_paragraphs": len(doc.paragraphs)}
                    except Exception as e:
                        extracted_text = f"[DOCX Parsing Error: {str(e)}]"
                else:
                    extracted_text = file_bytes.decode("utf-8", errors="ignore")

            # E-Mail Evidence (.eml)
            elif file_type == "email" or filename_lower.endswith(".eml") or filename_lower.endswith(".msg"):
                try:
                    msg = email.message_from_bytes(file_bytes)
                    subject = msg.get("subject", "")
                    sender = msg.get("from", "")
                    to = msg.get("to", "")
                    body = ""
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == "text/plain":
                                body += part.get_payload(decode=True).decode(errors="ignore")
                    else:
                        body = msg.get_payload(decode=True).decode(errors="ignore")
                    extracted_text = f"Subject: {subject}\nFrom: {sender}\nTo: {to}\n\n{body}"
                    metadata = {"subject": subject, "sender": sender, "to": to}
                except Exception as e:
                    extracted_text = f"[EML Parsing Error: {str(e)}]"

            # Text / Log / Default
            else:
                try:
                    extracted_text = file_bytes.decode("utf-8", errors="ignore")
                    metadata = {"line_count": len(extracted_text.splitlines())}
                except Exception as e:
                    extracted_text = f"[Binary/Text decode error: {str(e)}]"

        except Exception as e:
            extracted_text = f"[Extractor Global Error: {str(e)}]"

        return extracted_text, metadata

evidence_service = EvidenceService()
