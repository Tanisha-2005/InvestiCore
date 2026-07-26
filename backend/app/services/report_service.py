import os
from typing import Dict, Any, List
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class ReportService:
    def generate_pdf_report(self, output_path: str, case_data: Dict[str, Any], evidences: List[Dict[str, Any]], iocs: List[Dict[str, Any]]) -> str:
        doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        
        # Custom Styles
        title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor("#0f172a"), spaceAfter=12)
        h2_style = ParagraphStyle('SectionHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor("#1e293b"), spaceBefore=12, spaceAfter=6)
        body_style = ParagraphStyle('BodyTextCustom', parent=styles['BodyText'], fontSize=10, textColor=colors.HexColor("#334155"), spaceAfter=6)

        story = []

        # Header Title
        story.append(Paragraph(f"InvestiCore Investigation Report: {case_data.get('case_number')}", title_style))
        story.append(Paragraph(f"<b>Title:</b> {case_data.get('title')}", body_style))
        story.append(Paragraph(f"<b>Victim:</b> {case_data.get('victim_name', 'N/A')} | <b>Priority:</b> {case_data.get('priority', '').upper()} | <b>Risk Score:</b> {case_data.get('risk_score')}/100", body_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cbd5e1"), spaceBefore=8, spaceAfter=12))

        # Executive Summary
        story.append(Paragraph("Executive Summary", h2_style))
        story.append(Paragraph(case_data.get("summary_ai") or case_data.get("description") or "No executive summary available.", body_style))
        story.append(Spacer(1, 10))

        # Evidence Summary Table
        story.append(Paragraph("Uploaded Digital Evidence", h2_style))
        if evidences:
            ev_table_data = [["File Name", "Type", "Size (bytes)", "MD5 Hash"]]
            for ev in evidences:
                ev_table_data.append([
                    ev.get("file_name")[:25],
                    ev.get("file_type"),
                    str(ev.get("file_size")),
                    (ev.get("md5_hash") or "N/A")[:12] + "..."
                ])
            t = Table(ev_table_data, colWidths=[150, 80, 80, 150])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ]))
            story.append(t)
        else:
            story.append(Paragraph("No digital evidence uploaded yet.", body_style))

        story.append(Spacer(1, 14))

        # Indicators of Compromise Table
        story.append(Paragraph("Extracted Indicators of Compromise (IOCs)", h2_style))
        if iocs:
            ioc_table_data = [["Type", "Value", "Status", "Threat Score"]]
            for ioc in iocs:
                ioc_table_data.append([
                    ioc.get("ioc_type").upper(),
                    ioc.get("value")[:35],
                    ioc.get("status").upper(),
                    str(ioc.get("threat_score"))
                ])
            t_ioc = Table(ioc_table_data, colWidths=[80, 220, 90, 70])
            t_ioc.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ]))
            story.append(t_ioc)
        else:
            story.append(Paragraph("No IOCs extracted for this case.", body_style))

        doc.build(story)
        return output_path

    def generate_docx_report(self, output_path: str, case_data: Dict[str, Any], evidences: List[Dict[str, Any]], iocs: List[Dict[str, Any]]) -> str:
        doc = docx.Document()
        doc.add_heading(f"Investigation Report: {case_data.get('case_number')}", 0)

        p = doc.add_paragraph()
        p.add_run(f"Title: {case_data.get('title')}\n").bold = True
        p.add_run(f"Priority: {case_data.get('priority')} | Status: {case_data.get('status')}\n")
        p.add_run(f"Risk Score: {case_data.get('risk_score')}/100\n")

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(case_data.get("summary_ai") or case_data.get("description") or "N/A")

        doc.add_heading("Digital Evidence", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'File Name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Size'
        hdr_cells[3].text = 'MD5 Hash'
        for ev in evidences:
            row_cells = table.add_row().cells
            row_cells[0].text = ev.get('file_name', '')
            row_cells[1].text = ev.get('file_type', '')
            row_cells[2].text = str(ev.get('file_size', 0))
            row_cells[3].text = ev.get('md5_hash', 'N/A')

        doc.add_heading("Indicators of Compromise (IOCs)", level=1)
        table_ioc = doc.add_table(rows=1, cols=3)
        hdr_cells_ioc = table_ioc.rows[0].cells
        hdr_cells_ioc[0].text = 'Type'
        hdr_cells_ioc[1].text = 'Value'
        hdr_cells_ioc[2].text = 'Status'
        for ioc in iocs:
            row_cells = table_ioc.add_row().cells
            row_cells[0].text = ioc.get('ioc_type', '')
            row_cells[1].text = ioc.get('value', '')
            row_cells[2].text = ioc.get('status', '')

        doc.save(output_path)
        return output_path

    def generate_markdown_report(self, output_path: str, case_data: Dict[str, Any], evidences: List[Dict[str, Any]], iocs: List[Dict[str, Any]]) -> str:
        md_content = f"# InvestiCore Investigation Report: {case_data.get('case_number')}\n\n"
        md_content += f"**Title:** {case_data.get('title')}\n"
        md_content += f"**Victim:** {case_data.get('victim_name', 'N/A')} | **Priority:** {case_data.get('priority')} | **Risk Score:** {case_data.get('risk_score')}/100\n\n"
        md_content += f"## Executive Summary\n{case_data.get('summary_ai') or case_data.get('description') or 'N/A'}\n\n"

        md_content += "## Digital Evidence\n"
        md_content += "| File Name | Type | Size | MD5 |\n|---|---|---|---|\n"
        for ev in evidences:
            md_content += f"| {ev.get('file_name')} | {ev.get('file_type')} | {ev.get('file_size')} | {ev.get('md5_hash')} |\n"

        md_content += "\n## Extracted Indicators of Compromise (IOCs)\n"
        md_content += "| Type | Value | Status |\n|---|---|---|\n"
        for ioc in iocs:
            md_content += f"| {ioc.get('ioc_type')} | `{ioc.get('value')}` | {ioc.get('status')} |\n"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return output_path

report_service = ReportService()
